import base64
import io
import sys
import markdown
import yaml
import traceback
import uuid
from datetime import datetime
from docx import Document
from PIL import Image
from typing import List, Dict, Optional
from pathlib import Path

try:
    from .models import ConversionInput, ConversionResult, FileConversionRequest, FileConversionResponse, Question
    from .omni_client import OmniClient
    from .config import config
    from .paths import get_paths
except ImportError:
    from models import ConversionInput, ConversionResult, FileConversionRequest, FileConversionResponse, Question
    from omni_client import OmniClient
    from config import config
    from paths import get_paths

class FileParser:
    """
    File Parser which:
    - detect if there're embedded images in md and docx files
    - file extension consistency check & total size check
    - extract plain text content from txt, md and docx (without embedded images) files
    - convert pdf to base64 image strings in png format
    - convert images to base64 image strings in png format
    """

    SUPPORTED_FORMATS = {'.txt', '.md', '.docx', '.jpg', '.jpeg', '.png', '.pdf'}

    @staticmethod
    def extract_text_from_txt(base64_content: str) -> str:
        """Extract text from base64 content of .txt file"""
        content = base64.b64decode(base64_content)
        return content.decode()
    
    @staticmethod
    def extract_text_from_md(base64_content: str) -> str:
        """Extract text from base64 content of .md file"""
        content = base64.b64decode(base64_content)
        md_text = content.decode()
        html = markdown.markdown(md_text)
        import re
        plain_text = re.sub(r'<[^>]+>', '', html)
        return plain_text.strip()
    
    @staticmethod
    def extract_text_from_docx(base64_content: str) -> str:
        """Extract text from base64 content of .docx file"""
        content = base64.b64decode(base64_content)
        doc = Document(io.BytesIO(content))
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        return '\n'.join(text_parts)
    
    @staticmethod
    def convert_pdf_to_base64(base64_content: bytes) -> List[str]:
        """Convert PDF to list of PIL images"""
        content = base64.b64decode(base64_content)
        from pdf2image import convert_from_bytes
        images = convert_from_bytes(content, dpi=224, fmt="png")
        img_strs = []
        for image in images:
            buffered = io.BytesIO()
            image.save(buffered, format="PNG")
            img_str = base64.b64encode(buffered.getvalue()).decode('utf-8')
            img_strs.append(img_str)
        return img_strs
    
    @staticmethod
    def convert_image_to_base64(image: Image.Image) -> str:
        """Convert PIL Image to base64 string"""
        buffered = io.BytesIO()
        image.save(buffered, format="PNG")
        img_str = base64.b64encode(buffered.getvalue()).decode('utf-8')
        return img_str
    
    @classmethod
    def parse_files(cls, conversion_request: FileConversionRequest) -> ConversionInput:
        """
        Parse multiple files and return structured data.
        """
        filenames = conversion_request.filenames
        file_contents = conversion_request.file_contents

        if len(filenames) != len(file_contents):
            raise ValueError("Number of filenames and file contents must match")

        if not filenames:
            raise ValueError("No files provided")

        # 1. check if extension is supported
        extensions = []
        for filename in filenames:
            ext = Path(filename).suffix.lower() 
            if ext not in cls.SUPPORTED_FORMATS:
                raise ValueError(f"Unsupported file extension for '{filename}': {ext}. Supported formats: {cls.SUPPORTED_FORMATS}")
            extensions.append(ext)

        # 2. parse each file based on extension
        texts = []
        images = []
        try:
            for filename, ext, base64_content in zip(filenames, extensions, file_contents):
                if ext == '.txt':
                    texts.append(cls.extract_text_from_txt(base64_content))
                elif ext == '.md':
                    texts.append(cls.extract_text_from_md(base64_content))
                elif ext == '.docx':
                    texts.append(cls.extract_text_from_docx(base64_content))
                elif ext == '.pdf':
                    images.extend(cls.convert_pdf_to_base64(base64_content))
                else:
                    # Convert base64 to PIL Image first
                    content = base64.b64decode(base64_content)
                    image = Image.open(io.BytesIO(content))
                    images.append(cls.convert_image_to_base64(image))
        except Exception as e:
            stack_trace = ''.join(traceback.format_stack())
            raise ValueError(
                f"Text extraction or image conversion went wrong: {e}\n"
                f"Stack trace:\n{stack_trace}"
            )

        # 3. examine total image size
        size = sys.getsizeof(images)
        if size > 10 * 1024 * 1024:
            raise ValueError("File(s) too large. Recommend to upload files of total size <= 6MB")
        
        return ConversionInput(
            texts=texts,
            images=images
        )



class FileConverter:
    """File converter that handles the complete conversion pipeline"""

    def __init__(self):
        self._vl_client = None

    def get_vl_client(self):
        """Get or create OmniClient instance (lazy initialization)"""
        if self._vl_client is None:
            self._vl_client = OmniClient(config.get("models.vision_model", "qwen3-vl-plus"))
        return self._vl_client

    def _generate_yaml_content(self, original_filenames: List[str], questions: List[Question]) -> str:
        """Generate YAML content from extracted questions"""

        # Create exam structure
        exam_data = {
            "exam": {
                "title": f"Generated Exam from {', '.join(original_filenames)}",
                "description": f"Auto-generated exam from files: {', '.join(original_filenames)}",
                "section_instructions": {
                    "read_aloud": {
                        "text": "For read aloud questions, you should read displayed English sentences aloud.",
                        "tts": "For read aloud questions, you should read displayed English sentences aloud"
                    },
                    "multiple_choice": {
                        "text": "For multiple choice questions, read the question carefully and choose the best answer from A, B, C and D.",
                        "tts": "For multiple choice questions, read the question carefully and choose the best answer from A, B, C and D"
                    },
                    "quick_response": {
                        "text": "For quick response questions, you will hear a question and respond by speaking. Listen carefully and answer clearly.",
                        "tts": "For quick response questions, you will hear a question and respond by speaking. Listen carefully and answer clearly"
                    },
                    "translation": {
                        "text": "For translation questions, you will see a Chinese sentence and speak the English translation.",
                        "tts": "For translation questions, you will see a Chinese sentence and speak the English translation"
                    }
                },
                "questions": []
            }
        }

        # Add questions to exam data
        for question in questions:
            question_dict = {
                "id": question.id,
                "type": question.type,
                "text": question.text,
            }

            if question.options:
                question_dict["options"] = question.options

            if question.reference_answer:
                question_dict["reference_answer"] = question.reference_answer

            exam_data["exam"]["questions"].append(question_dict)

        # Generate YAML content
        return yaml.dump(
            exam_data,
            default_flow_style=False,
            allow_unicode=True,
            indent=2,
            width=100,
            sort_keys=False,
        )

    async def convert_files(self, request: FileConversionRequest) -> FileConversionResponse:
        """Convert files to exam YAML format"""

        # 1. Parse files
        conversion_input = FileParser.parse_files(request)

        # 2. Convert to questions with VLM
        conversion_result = await self.get_vl_client().convert_files_to_questions(conversion_input)
        if not conversion_result.success: # if not success, return directly
            return FileConversionResponse(
                success=False,
                message=conversion_result.message,
                extracted_questions=[]
            )

        # 3. Generate YAML content
        yaml_content = self._generate_yaml_content(request.filenames, conversion_result.extracted_questions)

        # 4. Generate output filename based on first filename
        base_filename = Path(request.filenames[0]).stem if request.filenames else "generated_exam"
        output_filename = f"{base_filename}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.yaml"

        # 5. Save YAML file
        paths = get_paths()
        exams_dir = paths.exams_dir
        output_path = exams_dir / output_filename

        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(yaml_content)
        except Exception as e:
            return FileConversionResponse(
                success=False,
                message=f"Failed to save YAML file: {str(e)}",
                extracted_questions=conversion_result.extracted_questions
            )

        return FileConversionResponse(
            success=True,
            message=conversion_result.message,
            extracted_questions=conversion_result.extracted_questions,
            yaml_output=yaml_content,
            output_filename=str(output_path)
        )

    async def rename_exam_file(self, old_name: str, new_name: str) -> dict:
        """Rename an exam file in the exams directory"""
        from pathlib import Path

        # Define exams directory path
        paths = get_paths()
        exams_dir = paths.exams_dir

        # Create paths
        old_path = exams_dir / old_name
        new_path = exams_dir / new_name

        # Check if old file exists
        if not old_path.exists():
            raise FileNotFoundError(f"File {old_name} not found in exams directory")

        # Check if new file already exists
        if new_path.exists():
            raise FileExistsError(f"File {new_name} already exists in exams directory")

        # Ensure new filename has .yaml extension
        if not new_name.endswith('.yaml'):
            new_name += '.yaml'
            new_path = exams_dir / new_name

        # Rename the file
        try:
            old_path.rename(new_path)
            return {
                "success": True,
                "message": f"Successfully renamed {old_name} to {new_name}",
                "new_filename": new_name
            }
        except Exception as e:
            raise Exception(f"Failed to rename file: {str(e)}")

    async def delete_exam_file(self, exam_filename: str) -> dict:
        """Delete an exam file from the exams directory"""
        from pathlib import Path

        # Define exams directory path
        paths = get_paths()
        exams_dir = paths.exams_dir

        # Create path
        exam_path = exams_dir / exam_filename

        # Check if file exists
        if not exam_path.exists():
            raise FileNotFoundError(f"File {exam_filename} not found in exams directory")

        # Security check: ensure file is in exams directory and has .yaml extension
        try:
            exam_path.resolve().relative_to(exams_dir.resolve())
        except ValueError:
            raise ValueError("Invalid file path - must be within exams directory")

        if not exam_filename.endswith('.yaml') and not exam_filename.endswith('.yml'):
            raise ValueError("Only .yaml and .yml files can be deleted")

        # Delete the file
        try:
            exam_path.unlink()
            return {
                "success": True,
                "message": f"Successfully deleted {exam_filename}"
            }
        except Exception as e:
            raise Exception(f"Failed to delete file: {str(e)}")

if __name__ == "__main__":
    import asyncio
    import base64
    from pathlib import Path

    async def test_convert_files():
        """Test the convert_files function directly"""
        print("Testing convert_files function directly...")

        # Initialize file converter
        converter = FileConverter()

        # Test with a single text file
        test_file_path = "backend/test_files/2018.docx"

        # Read and encode the test file
        with open(test_file_path, 'rb') as f:
            file_content = base64.b64encode(f.read()).decode('utf-8')

        # Create request
        request = FileConversionRequest(
            filenames=[test_file_path],
            file_contents=[file_content]
        )

        # Test the conversion
        result = await converter.convert_files(request)
        print(f"Success: {result.success}")
        print(f"Message: {result.message}")
        print(f"Number of questions: {len(result.extracted_questions)}")

        if result.output_filename:
            print(f"Output file: {result.output_filename}")

        # Show first few questions
        for i, q in enumerate(result.extracted_questions[:3]):
            print(f"\nQuestion {i+1}:")
            print(f"  ID: {q.id}")
            print(f"  Type: {q.type}")
            print(f"  Text: {q.text}")
            if q.options:
                print(f"  Options: {q.options}")
            if q.reference_answer:
                print(f"  Answer: {q.reference_answer}")

    # Run the test
    asyncio.run(test_convert_files())

